"""Word document operations using python-docx.

OOXML output is natively compatible with ONLYOFFICE, Microsoft Word,
LibreOffice Writer, and Google Docs.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .validation import (
    validate_path, validate_color, validate_align, sanitize_text,
    validate_choice, validate_bounded_int,
)

# Map of friendly style names to python-docx style names. python-docx loads
# the default Office template which ships with these names.
STYLE_MAP: dict[str, str] = {
    "heading1": "Heading 1",
    "heading2": "Heading 2",
    "heading3": "Heading 3",
    "heading4": "Heading 4",
    "heading5": "Heading 5",
    "heading6": "Heading 6",
    "title": "Title",
    "subtitle": "Subtitle",
    "body": "Normal",
    "normal": "Normal",
    "quote": "Quote",
    "code": "Intense Quote",
    "caption": "Caption",
    "list": "List Bullet",
    "bullet": "List Bullet",
    "bullet2": "List Bullet 2",
    "bullet3": "List Bullet 3",
    "numbered": "List Number",
    "numbered2": "List Number 2",
    "numbered3": "List Number 3",
}

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_style(paragraph, style_name: str, doc: Document) -> str | None:
    style_key = style_name.lower()
    docx_style = STYLE_MAP.get(style_key, style_name)
    try:
        paragraph.style = doc.styles[docx_style]
        return None
    except KeyError:
        available = [s.name for s in doc.styles if s.type == 1][:20]
        return (
            f"Style '{style_name}' not found in template (defaulting to Normal). "
            f"Available: {available}"
        )


def _apply_paragraph_format(para, block: dict) -> None:
    """Apply indentation, spacing, and keep_with_next from block dict."""
    pf = para.paragraph_format
    if block.get("left_indent") is not None:
        pf.left_indent = Pt(block["left_indent"])
    if block.get("right_indent") is not None:
        pf.right_indent = Pt(block["right_indent"])
    if block.get("first_line_indent") is not None:
        pf.first_line_indent = Pt(block["first_line_indent"])
    if block.get("space_before") is not None:
        pf.space_before = Pt(block["space_before"])
    if block.get("space_after") is not None:
        pf.space_after = Pt(block["space_after"])
    if block.get("line_spacing") is not None:
        val = block["line_spacing"]
        if isinstance(val, (int, float)) and val > 3:
            pf.line_spacing = Pt(val)
        else:
            pf.line_spacing = val
    if block.get("keep_with_next") is not None:
        pf.keep_with_next = block["keep_with_next"]


def _apply_run_format(para, block: dict) -> None:
    """Apply run-level formatting to all runs in a paragraph."""
    has_fmt = any(block.get(k) is not None for k in (
        "bold", "italic", "underline", "strikethrough",
        "font", "font_name", "size", "font_size", "color", "font_color",
    ))
    if not has_fmt:
        return
    for run in para.runs:
        if block.get("bold") is not None:
            run.bold = block["bold"]
        if block.get("italic") is not None:
            run.italic = block["italic"]
        if block.get("underline") is not None:
            run.underline = block["underline"]
        if block.get("strikethrough"):
            run.font.strike = True
        font_name = block.get("font") or block.get("font_name")
        if font_name:
            run.font.name = font_name
        font_size = block.get("size") or block.get("font_size")
        if font_size:
            run.font.size = _pt(font_size, "size")
        color_val = block.get("color") or block.get("font_color")
        if color_val:
            c = validate_color(color_val)
            run.font.color.rgb = RGBColor(
                int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
            )


def _set_cell_shading(cell, color_hex: str) -> None:
    """Apply background shading to a table cell via OOXML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _pt(val, name="size"):
    """Coerce a points value, raising an AI-friendly error on bad input."""
    try:
        return Pt(float(val))
    except (TypeError, ValueError):
        raise ValueError(f"{name} must be a number of points, got {val!r}.") from None


def _int_in(val, name, lo, hi):
    """Coerce an int and clamp to [lo, hi], with a friendly error on bad input."""
    try:
        return max(lo, min(int(val), hi))
    except (TypeError, ValueError):
        raise ValueError(f"{name} must be an integer {lo}-{hi}, got {val!r}.") from None


def _parse_rc(key, rows, cols):
    """Parse a 'row,col' cell key and bounds-check it against the table."""
    try:
        ri, ci = (int(x) for x in str(key).split(","))
    except (ValueError, TypeError):
        raise ValueError(
            f"cell key {key!r} must be 'row,col' with integers, e.g. '1,0' or '0,2'."
        ) from None
    if not (0 <= ri < rows and 0 <= ci < cols):
        raise ValueError(
            f"cell key '{key}' is out of range for a {rows} row x {cols} column table."
        )
    return ri, ci


def _add_block(doc: Document, item: Any) -> None:
    if isinstance(item, str):
        doc.add_paragraph(sanitize_text(item, "paragraph"))
        return
    if not isinstance(item, dict):
        raise ValueError(f"Unsupported paragraph item type: {type(item).__name__}")

    item_type = item.get("type", "paragraph")

    if item_type in ("paragraph", "text") or ("text" in item and item_type == "paragraph"):
        text = sanitize_text(item.get("text", ""), "paragraph text")
        style = item.get("style", "normal")
        para = doc.add_paragraph(text)
        _apply_style(para, style, doc)
        if item.get("align"):
            align = ALIGN_MAP.get(item["align"].lower())
            if align is not None:
                para.alignment = align
        has_fmt = any(item.get(k) for k in ("bold", "italic", "underline", "strikethrough", "color", "size", "font"))
        if has_fmt:
            for run in para.runs:
                if item.get("bold"):
                    run.bold = True
                if item.get("italic"):
                    run.italic = True
                if item.get("underline"):
                    run.underline = True
                if item.get("strikethrough"):
                    run.font.strike = True
                if item.get("font"):
                    run.font.name = item["font"]
                if item.get("size"):
                    run.font.size = _pt(item["size"], "size")
                if item.get("color"):
                    color = validate_color(item["color"])
                    run.font.color.rgb = RGBColor(
                        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                    )
        _apply_paragraph_format(para, item)

    elif item_type == "heading":
        text = sanitize_text(item.get("text", ""), "heading text")
        level = _int_in(item.get("level", 1), "heading level", 0, 9)
        if item.get("numbering_prefix"):
            prefix = sanitize_text(str(item["numbering_prefix"]), "numbering prefix")
            text = f"{prefix} {text}"
        para = doc.add_heading(text, level=level)
        if item.get("alignment"):
            para.alignment = ALIGN_MAP.get(validate_align(item["alignment"]))
        _apply_run_format(para, item)
        _apply_paragraph_format(para, item)

    elif item_type == "table":
        data = item.get("data", [])
        if not data:
            return
        if not isinstance(data, list) or not all(isinstance(r, list) for r in data):
            raise ValueError(
                "table 'data' must be a list of row lists.\n"
                'Example: {"type": "table", "data": [["H1", "H2"], ["a", "b"]]}'
            )
        rows = len(data)
        cols = max(len(row) for row in data)
        tbl = doc.add_table(rows=rows, cols=cols)
        tbl.style = item.get("style", "Light Grid Accent 1")
        for ri, row in enumerate(data):
            for ci, cell in enumerate(row):
                tbl.cell(ri, ci).text = "" if cell is None else sanitize_text(str(cell), "table cell")
        # Bold the header row if requested (default True).
        if item.get("header", True) and rows > 0:
            for cell in tbl.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        if item.get("col_widths"):
            for i, w in enumerate(item["col_widths"]):
                if i < len(tbl.columns):
                    try:
                        tbl.columns[i].width = Inches(float(w))
                    except (TypeError, ValueError):
                        raise ValueError(
                            f"col_widths[{i}] must be a number of inches, got {w!r}."
                        ) from None
        if item.get("header_shading") and rows > 0:
            hc = validate_color(item["header_shading"])
            for cell in tbl.rows[0].cells:
                _set_cell_shading(cell, hc)
        if item.get("cell_shading"):
            for key, color in item["cell_shading"].items():
                ri, ci = _parse_rc(key, rows, cols)
                _set_cell_shading(tbl.cell(ri, ci), validate_color(color))
        if item.get("cell_alignment"):
            for key, align in item["cell_alignment"].items():
                ri, ci = _parse_rc(key, rows, cols)
                a = ALIGN_MAP.get(validate_align(align))
                if a is not None:
                    for p in tbl.cell(ri, ci).paragraphs:
                        p.alignment = a

    elif item_type == "image":
        if not item.get("path"):
            raise ValueError(
                "image block requires a 'path' to an image file.\n"
                'Example: {"type": "image", "path": "/path/to/logo.png", "width_inches": 4}'
            )
        img_path = Path(item["path"]).expanduser().resolve()
        if not img_path.exists():
            raise ValueError(
                f"Image not found: {img_path}\nProvide a valid path to an existing image file."
            )
        width = item.get("width_inches")
        kwargs: dict[str, Any] = {}
        if width is not None:
            try:
                kwargs["width"] = Inches(float(width))
            except (TypeError, ValueError):
                raise ValueError(
                    f"image 'width_inches' must be a number, got {width!r}."
                ) from None
        doc.add_picture(str(img_path), **kwargs)

    elif item_type == "pagebreak":
        doc.add_page_break()

    elif item_type == "list":
        items = item.get("items", [])
        if not isinstance(items, list):
            raise ValueError(
                "list 'items' must be a list of strings or {text, level} dicts.\n"
                'Example: {"type": "list", "items": ["First", {"text": "Nested", "level": 1}]}'
            )
        ordered = item.get("ordered", False)
        bullet_char = item.get("bullet_char")
        bullet_styles = ["List Bullet", "List Bullet 2", "List Bullet 3"]
        number_styles = ["List Number", "List Number 2", "List Number 3"]

        for entry in items:
            if isinstance(entry, str):
                entry = {"text": entry, "level": 0}
            elif not isinstance(entry, dict):
                raise ValueError(
                    f"list item must be a string or a dict, got {type(entry).__name__}."
                )
            text = sanitize_text(str(entry.get("text", "")), "list item")
            level = _int_in(entry.get("level", 0), "list item level", 0, 5)

            if bullet_char and not ordered:
                para = doc.add_paragraph(f"{bullet_char} {text}")
                para.paragraph_format.left_indent = Pt(36 * (level + 1))
                if level > 0:
                    para.paragraph_format.first_line_indent = Pt(-18)
            else:
                styles = number_styles if ordered else bullet_styles
                style_name = styles[min(level, 2)]
                para = doc.add_paragraph(text)
                try:
                    para.style = doc.styles[style_name]
                except KeyError:
                    pass
                if level > 2:
                    para.paragraph_format.left_indent = Pt(36 * (level + 1))

            _apply_run_format(para, entry)
            _apply_paragraph_format(para, entry)

    else:
        raise ValueError(f"Unknown block type: {item_type!r}")


def create(
    path: str,
    paragraphs: list[Any],
    title: str | None = None,
    author: str | None = None,
    subject: str | None = None,
) -> str:
    """Create a .docx file at `path` with the given content blocks."""
    out = validate_path(path, expected_ext="docx", for_creation=True, operation="create")
    doc = Document()
    cp = doc.core_properties
    if title:
        cp.title = title
    if author:
        cp.author = author
    if subject:
        cp.subject = subject
    for item in paragraphs:
        _add_block(doc, item)
    doc.save(str(out))
    return str(out)


def read(path: str, include_tables: bool = True) -> str:
    """Return the textual content of a .docx file."""
    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="read")
    doc = Document(str(in_))
    lines: list[str] = [p.text for p in doc.paragraphs]
    if include_tables:
        for tbl in doc.tables:
            lines.append("")  # blank line before table
            for row in tbl.rows:
                cells = [c.text for c in row.cells]
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def append(path: str, paragraphs: list[Any]) -> str:
    """Append content blocks to an existing .docx file."""
    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="append")
    doc = Document(str(in_))
    for item in paragraphs:
        _add_block(doc, item)
    doc.save(str(in_))
    return str(in_)


def set_metadata(
    path: str,
    title: str | None = None,
    author: str | None = None,
    subject: str | None = None,
    keywords: str | None = None,
    comments: str | None = None,
) -> str:
    """Set the core document properties on an existing .docx file."""
    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="set_metadata")
    doc = Document(str(in_))
    cp = doc.core_properties
    if title is not None:
        cp.title = title
    if author is not None:
        cp.author = author
    if subject is not None:
        cp.subject = subject
    if keywords is not None:
        cp.keywords = keywords
    if comments is not None:
        cp.comments = comments
    doc.save(str(in_))
    return str(in_)


def set_page_setup(
    path: str,
    *,
    size: str | None = None,
    width_mm: float | None = None,
    height_mm: float | None = None,
    top_mm: float | None = None,
    bottom_mm: float | None = None,
    left_mm: float | None = None,
    right_mm: float | None = None,
    orientation: str | None = None,
) -> str:
    """Set page size, margins, and orientation on a .docx file.

    ``size`` accepts named sizes: letter, a4, a3, a5, legal, tabloid.
    Or set ``width_mm`` / ``height_mm`` directly for custom sizes.
    Margins are in millimetres. ``orientation`` is 'portrait' or 'landscape'.
    """
    from .validation import validate_page_size

    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="set_page_setup")
    doc = Document(str(in_))

    for section in doc.sections:
        if size:
            w_mm, h_mm = validate_page_size(size)
            section.page_width = Mm(w_mm)
            section.page_height = Mm(h_mm)
        if width_mm is not None:
            section.page_width = Mm(width_mm)
        if height_mm is not None:
            section.page_height = Mm(height_mm)
        if top_mm is not None:
            section.top_margin = Mm(top_mm)
        if bottom_mm is not None:
            section.bottom_margin = Mm(bottom_mm)
        if left_mm is not None:
            section.left_margin = Mm(left_mm)
        if right_mm is not None:
            section.right_margin = Mm(right_mm)
        if orientation is not None:
            from docx.enum.section import WD_ORIENT
            o = orientation.lower().strip()
            if o not in ("portrait", "landscape"):
                raise ValueError(
                    f"Invalid orientation: '{orientation}'.\n"
                    f"Valid values: 'portrait', 'landscape'"
                )
            section.orientation = (
                WD_ORIENT.LANDSCAPE if o == "landscape" else WD_ORIENT.PORTRAIT
            )
            if o == "landscape":
                w, h = section.page_width, section.page_height
                if w < h:
                    section.page_width, section.page_height = h, w
            else:
                w, h = section.page_width, section.page_height
                if w > h:
                    section.page_width, section.page_height = h, w

    doc.save(str(in_))
    return str(in_)


def insert_paragraph(path: str, paragraph_index: int, content: Any) -> str:
    """Insert a content block before ``paragraph_index`` (0-based).

    ``content`` uses the same schema as items in ``docx_create``'s paragraphs
    list: a plain string, or a dict with type/text/style/etc.
    """
    from .validation import validate_paragraph_index

    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="insert_paragraph")
    doc = Document(str(in_))
    validate_paragraph_index(doc, paragraph_index)
    target = doc.paragraphs[paragraph_index]._element
    para_count_before = len(doc.paragraphs)
    _add_block(doc, content)
    if len(doc.paragraphs) > para_count_before:
        new_element = doc.paragraphs[-1]._element
        new_element.getparent().remove(new_element)
        target.addprevious(new_element)
    doc.save(str(in_))
    return str(in_)


def delete_paragraph(path: str, paragraph_index: int) -> str:
    """Delete a paragraph by index from a .docx file."""
    from .validation import validate_paragraph_index

    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="delete_paragraph")
    doc = Document(str(in_))
    validate_paragraph_index(doc, paragraph_index)
    p_element = doc.paragraphs[paragraph_index]._element
    p_element.getparent().remove(p_element)
    doc.save(str(in_))
    return str(in_)


def edit_paragraph(
    path: str,
    paragraph_index: int,
    *,
    text: str | None = None,
    style: str | None = None,
    align: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    strikethrough: bool | None = None,
    font: str | None = None,
    size: int | None = None,
    color: str | None = None,
) -> dict:
    """Edit an existing paragraph's text and formatting in place.

    Only the parameters you provide are changed — everything else is
    preserved. Returns the paragraph's state after editing.
    """
    from .validation import validate_paragraph_index

    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="edit_paragraph")
    doc = Document(str(in_))
    validate_paragraph_index(doc, paragraph_index)
    para = doc.paragraphs[paragraph_index]

    if text is not None:
        clean = sanitize_text(text, "paragraph text")
        for run in para.runs[1:]:
            run._element.getparent().remove(run._element)
        if para.runs:
            para.runs[0].text = clean
        else:
            para.add_run(clean)

    if style is not None:
        _apply_style(para, style, doc)

    if align is not None:
        validate_align(align)
        alignment = ALIGN_MAP.get(align.lower())
        if alignment is not None:
            para.alignment = alignment

    for run in para.runs:
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if strikethrough is not None:
            run.font.strike = strikethrough
        if font is not None:
            run.font.name = sanitize_text(font, "font name")
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            c = validate_color(color)
            run.font.color.rgb = RGBColor(
                int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16)
            )

    doc.save(str(in_))

    result_runs = [_run_formatting(r) for r in para.runs]
    return {
        "path": str(in_),
        "paragraph_index": paragraph_index,
        "text": para.text,
        "style": para.style.name if para.style else "Normal",
        "alignment": _align_name(para.alignment),
        "runs": result_runs,
    }


def read_metadata(path: str) -> dict:
    """Read core document properties from a .docx file."""
    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="read_metadata")
    doc = Document(str(in_))
    cp = doc.core_properties
    return {
        "title": cp.title or "",
        "author": cp.author or "",
        "subject": cp.subject or "",
        "keywords": cp.keywords or "",
        "comments": cp.comments or "",
        "created": str(cp.created) if cp.created else "",
        "modified": str(cp.modified) if cp.modified else "",
        "last_modified_by": cp.last_modified_by or "",
        "revision": cp.revision,
        "category": cp.category or "",
    }


def _align_name(alignment) -> str:
    """Convert a WD_ALIGN_PARAGRAPH enum value to a friendly name."""
    _REVERSE_ALIGN = {v: k for k, v in ALIGN_MAP.items()}
    if alignment is None:
        return "left"
    return _REVERSE_ALIGN.get(alignment, str(alignment))


def _rgb_to_hex(rgb) -> str | None:
    """Convert an RGBColor to '#RRGGBB' string, or None."""
    if rgb is None:
        return None
    return f"#{str(rgb)}"


def _run_formatting(run) -> dict[str, Any]:
    """Extract formatting info from a single run."""
    fmt: dict[str, Any] = {"text": run.text}
    if run.bold:
        fmt["bold"] = True
    if run.italic:
        fmt["italic"] = True
    if run.underline:
        fmt["underline"] = True
    if run.font.strike:
        fmt["strikethrough"] = True
    if run.font.name:
        fmt["font"] = run.font.name
    if run.font.size:
        fmt["size_pt"] = run.font.size.pt
    color_rgb = run.font.color.rgb if run.font.color and run.font.color.rgb else None
    if color_rgb:
        fmt["color"] = _rgb_to_hex(color_rgb)
    return fmt


def get_formatting(path: str, paragraph_index: int) -> dict[str, Any]:
    """Inspect formatting of a specific paragraph and its runs."""
    from .validation import validate_paragraph_index

    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="get_formatting")
    doc = Document(str(in_))
    validate_paragraph_index(doc, paragraph_index)
    para = doc.paragraphs[paragraph_index]

    runs = [_run_formatting(r) for r in para.runs]
    style_name = para.style.name if para.style else "Normal"
    alignment = _align_name(para.alignment)
    validate_align(alignment)

    return {
        "paragraph_index": paragraph_index,
        "text": para.text,
        "style": style_name,
        "alignment": alignment,
        "runs": runs,
        "run_count": len(runs),
    }


def get_config(path: str) -> dict[str, Any]:
    """Detect full document configuration — page setup, styles, fonts, colors, sections."""
    from lxml import etree

    from .safety import safe_parse_xml
    from .validation import format_for_path

    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="get_config")
    doc = Document(str(in_))
    fmt = format_for_path(str(in_))

    sections_info = []
    for i, section in enumerate(doc.sections):
        sec = {
            "index": i,
            "page_width_mm": round(section.page_width / Mm(1), 1) if section.page_width else None,
            "page_height_mm": round(section.page_height / Mm(1), 1) if section.page_height else None,
            "top_margin_mm": round(section.top_margin / Mm(1), 1) if section.top_margin else None,
            "bottom_margin_mm": round(section.bottom_margin / Mm(1), 1) if section.bottom_margin else None,
            "left_margin_mm": round(section.left_margin / Mm(1), 1) if section.left_margin else None,
            "right_margin_mm": round(section.right_margin / Mm(1), 1) if section.right_margin else None,
        }
        orient = section.orientation
        sec["orientation"] = "landscape" if orient and orient == 1 else "portrait"
        sec["has_header"] = bool(section.header and section.header.paragraphs and
                                 any(p.text.strip() for p in section.header.paragraphs))
        sec["has_footer"] = bool(section.footer and section.footer.paragraphs and
                                 any(p.text.strip() for p in section.footer.paragraphs))
        if sec["has_header"]:
            sec["header_text"] = " ".join(p.text for p in section.header.paragraphs if p.text.strip())
        if sec["has_footer"]:
            sec["footer_text"] = " ".join(p.text for p in section.footer.paragraphs if p.text.strip())
        sections_info.append(sec)

    fonts_used: set[str] = set()
    colors_used: set[str] = set()
    styles_used: set[str] = set()

    for para in doc.paragraphs:
        if para.style:
            styles_used.add(para.style.name)
        for run in para.runs:
            if run.font.name:
                fonts_used.add(run.font.name)
            if run.font.color and run.font.color.rgb:
                hex_c = _rgb_to_hex(run.font.color.rgb)
                if hex_c:
                    colors_used.add(hex_c)

    styles_available = [
        {"name": s.name, "type": str(s.type)}
        for s in doc.styles
        if s.type == 1
    ][:30]

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = doc.element
    bg_el = root.find(f"{{{W_NS}}}background")
    background_color = None
    if bg_el is not None:
        background_color = bg_el.get(f"{{{W_NS}}}color")

    has_background_image = False
    for section in doc.sections:
        header_el = section.header._element
        for anchor in header_el.iter():
            if anchor.tag.endswith("}docPr") and anchor.get("name") == "PageBackground":
                has_background_image = True
                break

    has_watermark = False
    V_NS = "urn:schemas-microsoft-com:vml"
    for section in doc.sections:
        header_el = section.header._element
        for shape in header_el.iter(f"{{{V_NS}}}shape"):
            if shape.get("id") == "WatermarkShape":
                has_watermark = True
                break

    return {
        "format": fmt,
        "paragraph_count": len(doc.paragraphs),
        "section_count": len(doc.sections),
        "sections": sections_info,
        "fonts_used": sorted(fonts_used),
        "colors_used": sorted(colors_used),
        "styles_used": sorted(styles_used),
        "styles_available": styles_available,
        "background_color": background_color,
        "has_background_image": has_background_image,
        "has_watermark": has_watermark,
        "table_count": len(doc.tables),
        "metadata": read_metadata(path),
    }


def format_cell(
    path: str,
    table_index: int,
    row: int,
    col: int,
    *,
    text: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    color: str | None = None,
    size: float | None = None,
    font: str | None = None,
    align: str | None = None,
    vertical_align: str | None = None,
    shading: str | None = None,
) -> str:
    """Format a single table cell: text content, run styling (bold/italic/
    underline/color/size/font), horizontal + vertical alignment, and fill.

    ``table_index`` is the 0-based index into the document's tables;
    ``row``/``col`` are 0-based. This fills the gap where the create/append
    block API can only colour the header row — here any cell's text can be
    styled (e.g. white text on a dark fill)."""
    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="format_cell")
    try:
        table_index, row, col = int(table_index), int(row), int(col)
    except (TypeError, ValueError):
        raise ValueError(
            f"table_index, row and col must be integers, got "
            f"({table_index!r}, {row!r}, {col!r})."
        ) from None
    if size is not None and (not isinstance(size, (int, float)) or size <= 0):
        raise ValueError(f"size must be a positive number of points, got {size!r}.")
    va_norm = validate_choice(vertical_align, "vertical_align", ("top", "center", "middle", "bottom"))
    doc = Document(str(in_))
    if not doc.tables:
        raise ValueError("The document has no tables to format. Add a table first (docx_create/docx_append).")
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(
            f"table_index {table_index} out of range — the document has "
            f"{len(doc.tables)} table(s), so valid indices are 0..{len(doc.tables) - 1}."
        )
    tbl = doc.tables[table_index]
    nrows, ncols = len(tbl.rows), len(tbl.columns)
    if not (0 <= row < nrows and 0 <= col < ncols):
        raise ValueError(
            f"cell ({row},{col}) is out of range for table {table_index}, which is "
            f"{nrows} row(s) x {ncols} column(s) (valid row 0..{nrows - 1}, col 0..{ncols - 1})."
        )
    cell = tbl.cell(row, col)
    if text is not None:
        cell.text = sanitize_text(str(text), "cell text")
    if shading:
        _set_cell_shading(cell, validate_color(shading))
    if va_norm:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as _VA
        cell.vertical_alignment = {"top": _VA.TOP, "center": _VA.CENTER,
                                   "middle": _VA.CENTER, "bottom": _VA.BOTTOM}[va_norm]
    color_hex = validate_color(color) if color else None
    for para in cell.paragraphs:
        if align:
            a = ALIGN_MAP.get(validate_align(align))
            if a is not None:
                para.alignment = a
        for run in para.runs:
            if bold is not None:
                run.bold = bool(bold)
            if italic is not None:
                run.italic = bool(italic)
            if underline is not None:
                run.underline = bool(underline)
            if font:
                run.font.name = font
            if size:
                run.font.size = Pt(size)
            if color_hex:
                run.font.color.rgb = RGBColor(
                    int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
                )
    doc.save(str(in_))
    return str(in_)


def extract_images(path: str, out_dir: str | None = None) -> dict:
    """Extract every embedded image from a .docx into ``out_dir`` (defaults to
    ``<docname>_images`` next to the file). Returns a dict with the count,
    directory and saved file paths."""
    from . import safety
    in_ = validate_path(path, must_exist=True, expected_ext="docx", operation="extract_images")
    doc = Document(str(in_))
    target = Path(out_dir).expanduser().resolve() if out_dir else in_.parent / f"{in_.stem}_images"
    safety.check_path_safety(target)
    target.mkdir(parents=True, exist_ok=True)
    saved: list[str] = []
    for part in doc.part.package.iter_parts():
        name = str(part.partname)
        if name.startswith("/word/media/"):
            dest = target / Path(name).name
            dest.write_bytes(part.blob)
            saved.append(str(dest))
    return {"count": len(saved), "directory": str(target), "files": saved}
